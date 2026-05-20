#!/usr/bin/env python3
"""
update_word.py — Sync Excel task data into 版本发布更新说明-管维.docx

Reads excel/ (auto-discovered): 新增功能, 修复缺陷, 遗留问题, 版本信息.
Updates the 管维平台 section (chapter 5):
  版本信息 table → B2 date + B1 version
  功能更新 table → 新增功能 sheet rows
  已修复问题 table → 修复缺陷 sheet rows
  已知问题与限制 → 遗留问题 sheet rows
  升级说明 paragraph → auto-generated summary

Usage:
    python update_word.py                # auto-discover files
    python update_word.py --inspect      # print doc structure only
    python update_word.py --excel path --word path
"""

import os
import re
import sys
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
EXCEL_DIR  = os.path.join(SCRIPT_DIR, "excel")
WORD_DIR   = os.path.join(SCRIPT_DIR, "word")

SHEET_FEATURE = "\u65b0\u589e\u529f\u80fd "   # 新增功能
SHEET_BUG     = "\u4fee\u590d\u7f3a\u9677"    # 修复缺陷
SHEET_LEGACY  = "\u9057\u7559\u95ee\u9898"    # 遗留问题
SHEET_VERSION = "\u7248\u672c\u4fe1\u606f"    # 版本信息

_W      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TAG_P  = f"{{{_W}}}p"
_TAG_T  = f"{{{_W}}}tbl"
_TAG_X  = f"{{{_W}}}t"


# ── helpers ───────────────────────────────────────────────────────────────────

def _require(pkg, install_name):
    try:
        return __import__(pkg)
    except ImportError:
        print(f"ERROR: {pkg} not installed.  Run: pip install {install_name}",
              file=sys.stderr)
        sys.exit(1)


def _resolve(directory, ext, explicit=None, keyword=None):
    if explicit:
        return explicit
    found = [f for f in os.listdir(directory)
             if f.lower().endswith(ext) and not f.startswith("~$")]
    if not found:
        print(f"ERROR: no {ext} in {directory}", file=sys.stderr)
        sys.exit(1)
    if len(found) == 1:
        return os.path.join(directory, found[0])
    if keyword:
        kw = [f for f in found if keyword in f]
        if len(kw) == 1:
            return os.path.join(directory, kw[0])
    print(f"Multiple {ext} in {directory} — use --excel/--word to specify:",
          file=sys.stderr)
    for f in found:
        print(f"  {f}", file=sys.stderr)
    sys.exit(1)


def _read_sheet_rows(wb, sheet_name):
    """Return data rows (skip header) as list of dicts keyed by header text."""
    ws = wb[sheet_name]
    headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
    rows = []
    for r in range(2, ws.max_row + 1):
        vals = [ws.cell(r, c).value for c in range(1, len(headers) + 1)]
        if any(v is not None for v in vals) and vals[0] is not None:
            rows.append({h: v for h, v in zip(headers, vals)})
    return rows


def _para_text(elem):
    return "".join(t.text or "" for t in elem.iter(_TAG_X))


def _table_after_section(doc, section_heading, sub_heading):
    """
    Find the first <w:tbl> after a <w:p> containing sub_heading, but only
    within the section that starts at the paragraph containing section_heading.
    This avoids matching wrong-section tables (e.g. RTOS vs 管维).
    """
    from docx.table import Table
    in_section = False
    for child in doc.element.body:
        if child.tag == _TAG_P:
            txt = _para_text(child)
            if section_heading in txt:
                in_section = True
                continue
            if in_section and sub_heading in txt:
                # find next table
                idx = list(doc.element.body).index(child)
                for later in list(doc.element.body)[idx + 1:]:
                    if later.tag == _TAG_T:
                        return Table(later, doc)
                    if later.tag == _TAG_P:
                        t2 = _para_text(later).strip()
                        if t2:
                            break  # non-empty paragraph means table not immediately following
                return None
        if child.tag == _TAG_P and in_section:
            # next major heading ends this section
            pass
    return None


def _clear_data_rows(table):
    for i in range(len(table.rows) - 1, 0, -1):
        tr = table.rows[i]._tr
        tr.getparent().remove(tr)


def _write_cell(cell, text):
    """Write text into a cell, reusing first run for formatting."""
    para = cell.paragraphs[0]
    s = str(text) if text is not None else ""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = s
    else:
        para.add_run(s)


def _fill_table(table, data_rows, col_title=1):
    """
    Clear data rows and repopulate: col0=seq, col{col_title}=标题, col2=描述.
    Preserves header; clears all data rows then rebuilds.
    """
    _clear_data_rows(table)
    for _ in range(len(data_rows)):
        table.add_row()
    for seq, row_dict in enumerate(data_rows, 1):
        r = table.rows[seq]
        _write_cell(r.cells[0], seq)
        if col_title < len(r.cells):
            _write_cell(r.cells[col_title], row_dict.get("\u6807\u9898", ""))
        if 2 < len(r.cells) and col_title != 2:
            _write_cell(r.cells[2], row_dict.get("\u63cf\u8ff0", ""))
    print(f"  [word] table: {len(data_rows)} rows filled", flush=True)


def _update_upgrade_para(doc, section_heading, feat_rows, bug_rows, legacy_rows, version, date):
    """
    Find 升级说明 paragraph inside the given section, then replace the
    next body paragraph's text with an auto-generated summary.
    """
    paras = doc.paragraphs
    in_section = False
    for i, p in enumerate(paras):
        if section_heading in p.text:
            in_section = True
            continue
        if in_section and "\u5347\u7ea7\u8bf4\u660e" in p.text:  # 升级说明
            # Find the next non-empty paragraph
            for j in range(i + 1, len(paras)):
                if paras[j].text.strip():
                    target = paras[j]
                    lines = [
                        f"\u672c\u7248\u672c\uff08{version}\uff0c{date}\uff09\u4e3b\u8981\u53d8\u66f4\u5982\u4e0b\uff1a"
                    ]
                    if feat_rows:
                        lines.append(f"\u65b0\u589e\u529f\u80fd\uff08\u5171 {len(feat_rows)} \u9879\uff09\uff1a")
                        for r in feat_rows:
                            lines.append(f"  \u00b7 {r.get('\u6807\u9898', '')}")
                    if bug_rows:
                        lines.append(f"\u5df2\u4fee\u590d\u95ee\u9898\uff08\u5171 {len(bug_rows)} \u9879\uff09\uff1a")
                        for r in bug_rows:
                            lines.append(f"  \u00b7 {r.get('\u6807\u9898', '')}")
                    if legacy_rows:
                        lines.append(f"\u9057\u7559\u5df2\u77e5\u95ee\u9898\uff08\u5171 {len(legacy_rows)} \u9879\uff09\uff1a")
                        for r in legacy_rows:
                            lines.append(f"  \u00b7 {r.get('\u6800\u9898', '')}")

                    summary = "\n".join(lines)
                    for run in target.runs:
                        run.text = ""
                    if target.runs:
                        target.runs[0].text = summary
                    else:
                        target.add_run(summary)
                    print(f"  [word] \u5347\u7ea7\u8bf4\u660e: updated", flush=True)
                    return
    print(f"  [word] WARNING: \u5347\u7ea7\u8bf4\u660e in [{section_heading}] not found",
          flush=True)


# ── inspect mode ──────────────────────────────────────────────────────────────

def inspect_doc(word_path):
    from docx import Document
    doc = Document(word_path)
    print(f"\n=== PARAGRAPHS ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"  [{i:3d}] style={p.style.name!r:20s} | {p.text[:70]!r}")
    print(f"\n=== TABLES ===")
    for ti, t in enumerate(doc.tables):
        print(f"\n  -- Table {ti} ({len(t.rows)}r x {len(t.columns)}c) --")
        for ri, row in enumerate(t.rows[:5]):
            print(f"    row[{ri}]: {[c.text.strip()[:40] for c in row.cells]}")
        if len(t.rows) > 5:
            print(f"    ... ({len(t.rows)-5} more rows)")


# ── main update ───────────────────────────────────────────────────────────────

def update_word(excel_path=None, word_path=None, descriptions=None):
    _require("openpyxl", "openpyxl")
    _require("docx", "python-docx")
    import openpyxl
    from docx import Document

    excel_path = _resolve(EXCEL_DIR, ".xlsx", excel_path)
    word_path  = _resolve(WORD_DIR,  ".docx", word_path, keyword="管维")

    # Load descriptions sidecar if not passed in
    if descriptions is None:
        sidecar = os.path.join(EXCEL_DIR, "task_descriptions.json")
        if os.path.exists(sidecar):
            import json
            with open(sidecar, encoding="utf-8") as f:
                descriptions = json.load(f)

    wb = openpyxl.load_workbook(excel_path)

    # Version info
    vs       = wb[SHEET_VERSION]
    version  = str(vs.cell(1, 2).value or "").strip()
    raw_date = str(vs.cell(2, 2).value or "").strip()
    if re.fullmatch(r"\d{8}", raw_date):
        dt = datetime.strptime(raw_date, "%Y%m%d")
        build_date = f"{dt.year}-{dt.month}-{dt.day}"
    else:
        build_date = raw_date

    feat_rows   = _read_sheet_rows(wb, SHEET_FEATURE)
    bug_rows    = _read_sheet_rows(wb, SHEET_BUG)
    legacy_rows = _read_sheet_rows(wb, SHEET_LEGACY)

    # Inject descriptions into row dicts
    if descriptions:
        for row_list in (feat_rows, bug_rows, legacy_rows):
            for rd in row_list:
                tid = str(rd.get("编号", "") or "")
                if tid in descriptions:
                    rd["描述"] = descriptions[tid]

    doc = Document(word_path)
    SECTION = "\u7ba1\u7ef4\u5e73\u53f0\u7248\u672c\u66f4\u65b0\u8bf4\u660e"  # 管维平台版本更新说明

    # 1. Version table (Table 0): find row with 管维
    for table in doc.tables:
        if table.rows and "\u4ea7\u54c1" in table.rows[0].cells[0].text:
            for row in table.rows[1:]:
                if "\u7ba1\u7ef4" in row.cells[0].text:
                    _write_cell(row.cells[1], build_date)
                    _write_cell(row.cells[2], version)
                    print(f"  [word] version: \u7ba1\u7ef4\u5e73\u53f0 \u2192 {build_date}  {version}",
                          flush=True)
                    break
            break

    # 2-4. Section tables
    for sub_heading, rows, label in [
        ("\u529f\u80fd\u66f4\u65b0",       feat_rows,   "5.1 \u529f\u80fd\u66f4\u65b0"),
        ("\u5df2\u4fee\u590d\u95ee\u9898", bug_rows,    "5.2 \u5df2\u4fee\u590d\u95ee\u9898"),
        ("\u5df2\u77e5\u95ee\u9898",       legacy_rows, "5.3 \u5df2\u77e5\u95ee\u9898\u4e0e\u9650\u5236"),
    ]:
        t = _table_after_section(doc, SECTION, sub_heading)
        if t:
            _fill_table(t, rows)
            print(f"  [word] {label}: {len(rows)} rows", flush=True)
        else:
            print(f"  [word] WARNING: '{sub_heading}' table in {SECTION!r} not found",
                  flush=True)

    # 5. 升级说明 paragraph
    _update_upgrade_para(doc, SECTION, feat_rows, bug_rows, legacy_rows, version, build_date)

    try:
        doc.save(word_path)
        print(f"  [word] saved: {word_path}", flush=True)
    except PermissionError:
        print(f"  [word] ERROR: file is open in Word — close it and re-run",
              file=sys.stderr)
        sys.exit(1)


def main():
    args = sys.argv[1:]
    excel_path = word_path = None

    if "--inspect" in args:
        _require("docx", "python-docx")
        inspect_doc(_resolve(WORD_DIR, ".docx", keyword="\u7ba1\u7ef4"))
        return

    if "--excel" in args:
        i = args.index("--excel")
        excel_path = args[i + 1]
    if "--word" in args:
        i = args.index("--word")
        word_path = args[i + 1]

    update_word(excel_path, word_path)


if __name__ == "__main__":
    main()
